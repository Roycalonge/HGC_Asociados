from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento
doc = Document()

# Configuración de márgenes
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Estilo base
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Portada
title = doc.add_paragraph("DOSSIER ACADÉMICO E INNOVADOR\nHGC ASOCIADOS")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.runs[0]
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(212, 175, 55)

subtitle = doc.add_paragraph("Versión Institucional — Octubre 2025\nProyecto Académico – Contaduría Pública, Primer Semestre")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

lema = doc.add_paragraph("“Compromiso, ética y resultado.”")
lema.alignment = WD_ALIGN_PARAGRAPH.CENTER
lema.runs[0].italic = True
lema.runs[0].font.size = Pt(13)

doc.add_page_break()

# Función para crear secciones con línea dorada
def add_section(title_text, body_text):
    title_p = doc.add_paragraph(title_text)
    title_run = title_p.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(212, 175, 55)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body_p = doc.add_paragraph(body_text)
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    line = doc.add_paragraph("──────────────────────────────────────────────")
    line.runs[0].font.color.rgb = RGBColor(212, 175, 55)

# Secciones principales
add_section("📘 2. Presentación del Proyecto",
"""HGC Asociados es una iniciativa académica con enfoque académico, innovador, ético y digital.
Su propósito es integrar la teoría universitaria con la práctica aplicada, utilizando herramientas tecnológicas y estructuras organizativas reales.""")

add_section("⚙️ 3. Marco Ético y Normativo",
"""Basado en la Ley 43 de 1990, promueve integridad, confidencialidad y responsabilidad profesional.
El código guía la práctica y la comunicación dentro del grupo, asegurando profesionalismo, respeto y coherencia institucional.""")

add_section("🧩 4. Estructura Organizacional Actualizada",
"""DIRECTOR GENERAL (Contador Profesional Senior)
│
├── COORDINADOR ACADÉMICO
│   ├── Gestor de Contenidos Digitales
│   └── Tutor de Métodos (Sesiones por Telegram)
│
├── LÍDER DE PROYECTOS
│   ├── Planificador (Gestión vía Telegram)
│   └── Supervisor de Calidad (Reportes automáticos y Drive)
│
├── ESPECIALISTA TECNOLÓGICO (Ingeniería de Sistemas)
│   ├── Administrador de Plataformas (Telegram, Web, Drive, Bots)
│   ├── Desarrollador de Automatizaciones y Seguridad
│   └── Responsable de Innovación Digital
│
└── AUDITOR INTERNO
    ├── Garante Ético Digital
    └── Evaluador de Procesos Comunicativos""")

add_section("🧾 5. Manual Operativo (Resumen)",
"""Define los procesos de planeación, control, registro y evaluación del grupo.
Procesos clave: planificación semanal por Telegram, archivo en Drive, comunicación formal y evaluación de avances.""")

add_section("💬 6. Protocolos de Comunicación",
"""Incluyen comunicación interna (grupos, roles, bots) y externa (docentes, alianzas, sitio web).
Toda comunicación debe reflejar el mismo nivel de respeto y precisión que una firma contable real.""")

add_section("📚 7. Integración Académica por Asignaturas",
"""Economía: evaluación de recursos humanos y digitales.
Administración: estructura organizacional.
Contabilidad: Ley 43, ética y control documental.
Comunicación: redacción institucional.
Filosofía: pensamiento crítico.
Democracia: estructura normativa del grupo.""")

add_section("🌐 8. Ecosistema Digital",
"""Telegram, Google Drive, Página Web HGC, ChatGPT, Canva y Docs.
Cada plataforma tiene un propósito específico en la gestión y documentación institucional.""")

add_section("🚀 9. Proyección Interdisciplinaria y Estudiantil",
"""Incluye participación de Ingeniería de Sistemas, Derecho, Comunicación y Administración.
Creación de HGC Academy y desarrollo de herramientas digitales académicas.""")

add_section("📈 10. Impacto Académico",
"""Transforma teoría en práctica real desde el primer semestre, fomenta liderazgo, ética y gestión interdisciplinaria mediante herramientas digitales.""")

add_section("✍️ 11. Conclusión General",
"""HGC Asociados demuestra que la contaduría puede aplicarse de forma integral, interdisciplinaria y tecnológica.
Este dossier representa la madurez institucional de una firma académica en formación.""")

# Guardar documento
doc.save("Dossier_Academico_HGC_Asociados_2025.docx")
print("✅ Documento generado exitosamente: Dossier_Academico_HGC_Asociados_2025.docx")
